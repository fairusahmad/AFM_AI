
"""Translate the English AFM Operation Manual to Chinese while preserving DOCX formatting."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def translate_text(text):
    """Apply translation mappings to the text."""
    if not text or not text.strip():
        return text
    
    # Title page
    translations = {
        "AFM Hysteresis Simulation System": "AFM 迟滞效应模拟系统",
        "Operation Manual": "操作手册",
        
        # TOC
        "Table of Contents": "目录",
        "1. System Overview": "1. 系统概述",
        "2. Getting Started": "2. 入门指南",
        "3. Interface Layout": "3. 界面布局",
        "4. Navigation and Viewport Control": "4. 导航与视口控制",
        "5. Hysteresis Simulation and Compensation": "5. 迟滞模拟与补偿",
        "6. Optics and Focus Control": "6. 光学与对焦控制",
        "7. Sample Relocation Workflow": "7. 样品重定位工作流程",
        "8. AI-Assisted Features": "8. AI 辅助功能",
        "9. Data Collection and Export": "9. 数据采集与导出",
        "10. Keyboard and Mouse Reference": "10. 键盘与鼠标参考",
        "11. Troubleshooting": "11. 故障排除",
        
        # Section 1
        "Core Capabilities": "核心功能",
        "Architecture at a Glance": "系统架构概览",
        
        # Section 2
        "2.1 System Requirements": "2.1 系统要求",
        "2.2 Installation": "2.2 安装",
        "2.3 Launching the Application": "2.3 启动应用程序",
        "2.4 Default Sample Image": "2.4 默认样品图像",
        
        # Section 3
        "Viewport": "视口",
        "Navigation Dock": "导航面板",
        "Motion Dock": "运动面板",
        "Status Dock": "状态面板",
        "Relocation Trace Dock": "重定位轨迹面板",
        "Relocation Dock": "重定位面板",
        "Utility Dock": "工具面板",
        "Dock Layout Persistence": "面板布局持久化",
        
        # Section 4
        "4.1 Stage Movement": "4.1 载物台移动",
        "4.2 Zoom Control": "4.2 变焦控制",
        "4.3 Right-Click Context Menu": "4.3 右键上下文菜单",
        "4.4 Scale Bar": "4.4 比例尺",
        
        # Section 5
        "5.1 PI Hysteresis Model": "5.1 PI 迟滞模型",
        "5.2 PI Compensation Mode": "5.2 PI 补偿模式",
        "5.3 AI Inverse-Model Compensation": "5.3 AI 逆向模型补偿",
        "5.4 Auto Scan": "5.4 自动扫描",
        
        # Section 6
        "6.1 Optical Model": "6.1 光学模型",
        "6.2 Focus Adjustment": "6.2 对焦调整",
        "6.3 Stage Tilt": "6.3 载物台倾斜",
        "6.4 HUD Modes": "6.4 HUD 叠加模式",
        
        # Section 7
        "Step 1 — Save Region": "步骤 1 — 保存区域",
        "Step 2 — Remount": "步骤 2 — 重新装载",
        "Step 3 — Pick Origin": "步骤 3 — 选择原点",
        "Step 4 — Find Origin": "步骤 4 — 查找原点",
        "Step 5 — Recover Site (Coarse-to-Fine Relocation)": "步骤 5 — 恢复站点（从粗到精的重定位）",
        "Step 6 — Verify Tip": "步骤 6 — 验证探针位置",
        "Recommended Workflow": "推荐工作流程",
        
        # Section 8
        "8.1 Artefact Detection (YOLOv8)": "8.1 伪影检测（YOLOv8）",
        "8.2 Phase-2 Machine Learning Models": "8.2 第二阶段机器学习模型",
        "8.3 DeepSeek LLM Agent": "8.3 DeepSeek 大语言模型代理",
        "8.4 Training the AI Models": "8.4 训练 AI 模型",
        
        # Section 9
        "9.1 Automatic Data Collection": "9.1 自动数据采集",
        "9.2 Trajectory Export": "9.2 轨迹导出",
        
        # Section 10
        "10.1 Mouse Actions": "10.1 鼠标操作",
        "10.2 Viewport Arrow Controls": "10.2 视口方向箭头控制",
        "10.3 Button Reference": "10.3 按钮参考",
        
        # Section 11
    }
    
    # Apply exact match translations first
    if text in translations:
        return translations[text]
    
    # Check if text starts with a known heading pattern
    for en, zh in translations.items():
        if text.startswith(en + " "):
            return zh + text[len(en):]
    
    return text

def translate_paragraph(doc, para, translated_sentences):
    """Translate a paragraph and apply formatting."""
    text = para.text.strip()
    if not text:
        return text
    
    # Check exact translations
    exact_translations = {
        "AFM Hysteresis Simulation System": "AFM 迟滞效应模拟系统",
        "Operation Manual": "操作手册",
        "Table of Contents": "目录",
        "1. System Overview": "1. 系统概述",
        "2. Getting Started": "2. 入门指南",
        "3. Interface Layout": "3. 界面布局",
        "4. Navigation and Viewport Control": "4. 导航与视口控制",
        "5. Hysteresis Simulation and Compensation": "5. 迟滞模拟与补偿",
        "6. Optics and Focus Control": "6. 光学与对焦控制",
        "7. Sample Relocation Workflow": "7. 样品重定位工作流程",
        "8. AI-Assisted Features": "8. AI 辅助功能",
        "9. Data Collection and Export": "9. 数据采集与导出",
        "10. Keyboard and Mouse Reference": "10. 键盘与鼠标参考",
        "11. Troubleshooting": "11. 故障排除",
        "Core Capabilities": "核心功能",
        "Architecture at a Glance": "系统架构概览",
        "2.1 System Requirements": "2.1 系统要求",
        "2.2 Installation": "2.2 安装",
        "2.3 Launching the Application": "2.3 启动应用程序",
        "2.4 Default Sample Image": "2.4 默认样品图像",
        "Viewport": "视口",
        "Navigation Dock": "导航面板",
        "Motion Dock": "运动面板",
        "Status Dock": "状态面板",
        "Relocation Trace Dock": "重定位轨迹面板",
        "Relocation Dock": "重定位面板",
        "Utility Dock": "工具面板",
        "Dock Layout Persistence": "面板布局持久化",
        "4.1 Stage Movement": "4.1 载物台移动",
        "4.2 Zoom Control": "4.2 变焦控制",
        "4.3 Right-Click Context Menu": "4.3 右键上下文菜单",
        "4.4 Scale Bar": "4.4 比例尺",
        "5.1 PI Hysteresis Model": "5.1 PI 迟滞模型",
        "5.2 PI Compensation Mode": "5.2 PI 补偿模式",
        "5.3 AI Inverse-Model Compensation": "5.3 AI 逆向模型补偿",
        "5.4 Auto Scan": "5.4 自动扫描",
        "6.1 Optical Model": "6.1 光学模型",
        "6.2 Focus Adjustment": "6.2 对焦调整",
        "6.3 Stage Tilt": "6.3 载物台倾斜",
        "6.4 HUD Modes": "6.4 HUD 叠加模式",
        "Step 1 — Save Region": "步骤 1 — 保存区域",
        "Step 2 — Remount": "步骤 2 — 重新装载",
        "Step 3 — Pick Origin": "步骤 3 — 选择原点",
        "Step 4 — Find Origin": "步骤 4 — 查找原点",
        "Step 5 — Recover Site (Coarse-to-Fine Relocation)": "步骤 5 — 恢复站点（粗到精重定位）",
        "Step 6 — Verify Tip": "步骤 6 — 验证探针位置",
        "Recommended Workflow": "推荐工作流程",
        "8.1 Artefact Detection (YOLOv8)": "8.1 伪影检测（YOLOv8）",
        "8.2 Phase-2 Machine Learning Models": "8.2 第二阶段机器学习模型",
        "Same-Site Classifier": "同站点分类器",
        "Remount Transform Predictor": "重新装载变换预测器",
        "Low-Mag Embedding Index": "低倍镜嵌入索引",
        "8.3 DeepSeek LLM Agent": "8.3 DeepSeek 大语言模型代理",
        "8.4 Training the AI Models": "8.4 训练 AI 模型",
        "9.1 Automatic Data Collection": "9.1 自动数据采集",
        "9.2 Trajectory Export": "9.2 轨迹导出",
        "10.1 Mouse Actions": "10.1 鼠标操作",
        "10.2 Viewport Arrow Controls": "10.2 视口方向箭头控制",
        "10.3 Button Reference": "10.3 按钮参考",
        "— End of Manual —": "— 手册完 —",
    }
    
    if text in exact_translations:
        return exact_translations[text]
    
    # Check if it starts with a known pattern
    for en, zh in exact_translations.items():
        if text.startswith(en + ": "):
            return zh + ": " + text[len(en) + 2:]
        if text.startswith(en + " \u2014 "):
            return zh + " \u2014 " + text[len(en) + 3:]
    
    return None


# Read the English document
doc = Document('AFM_Hysteresis_Simulation_Operation_Manual.docx')

# We'll rebuild the document with translations
new_doc = Document()

# Copy styles
style = new_doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    heading_style = new_doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0x27, 0x47, 0x6A)

# Translation data structure for full paragraph translations
# We'll do a comprehensive manual translation of every paragraph

translations_map = {
    # Title page
    "AFM Hysteresis Simulation System": "AFM 迟滞效应模拟系统",
    "Operation Manual": "操作手册",
    
    # TOC
    "Table of Contents": "目录",
    "1. System Overview": "1. 系统概述",
    "2. Getting Started": "2. 入门指南",
    "3. Interface Layout": "3. 界面布局",
    "4. Navigation and Viewport Control": "4. 导航与视口控制",
    "5. Hysteresis Simulation and Compensation": "5. 迟滞模拟与补偿",
    "6. Optics and Focus Control": "6. 光学与对焦控制",
    "7. Sample Relocation Workflow": "7. 样品重定位工作流程",
    "8. AI-Assisted Features": "8. AI 辅助功能",
    "9. Data Collection and Export": "9. 数据采集与导出",
    "10. Keyboard and Mouse Reference": "10. 键盘与鼠标参考",
    "11. Troubleshooting": "11. 故障排除",
    
    # === SECTION 1 ===
    "Core Capabilities": "核心功能",
    "Architecture at a Glance": "系统架构概览",
    
    # === SECTION 2 ===
    "2.1 System Requirements": "2.1 系统要求",
    "2.2 Installation": "2.2 安装",
    "2.3 Launching the Application": "2.3 启动应用程序",
    "2.4 Default Sample Image": "2.4 默认样品图像",
    
    # === SECTION 3 ===
    "Viewport": "视口（Viewport）",
    "Navigation Dock": "导航面板（Navigation Dock）",
    "Motion Dock": "运动面板（Motion Dock）",
    "Status Dock": "状态面板（Status Dock）",
    "Relocation Trace Dock": "重定位轨迹面板（Relocation Trace Dock）",
    "Relocation Dock": "重定位面板（Relocation Dock）",
    "Utility Dock": "工具面板（Utility Dock）",
    "Dock Layout Persistence": "面板布局持久化",
    
    # === SECTION 4 ===
    "4.1 Stage Movement": "4.1 载物台移动",
    "4.2 Zoom Control": "4.2 变焦控制",
    "4.3 Right-Click Context Menu": "4.3 右键上下文菜单",
    "4.4 Scale Bar": "4.4 比例尺",
    
    # === SECTION 5 ===
    "5.1 PI Hysteresis Model": "5.1 PI 迟滞模型",
    "5.2 PI Compensation Mode": "5.2 PI 补偿模式",
    "5.3 AI Inverse-Model Compensation": "5.3 AI 逆向模型补偿",
    "5.4 Auto Scan": "5.4 自动扫描",
    
    # === SECTION 6 ===
    "6.1 Optical Model": "6.1 光学模型",
    "6.2 Focus Adjustment": "6.2 对焦调整",
    "6.3 Stage Tilt": "6.3 载物台倾斜",
    "6.4 HUD Modes": "6.4 HUD 叠加模式",
    
    # === SECTION 7 ===
    "Step 1 — Save Region": "步骤 1 — 保存区域",
    "Step 2 — Remount": "步骤 2 — 重新装载",
    "Step 3 — Pick Origin": "步骤 3 — 选择原点",
    "Step 4 — Find Origin": "步骤 4 — 查找原点",
    "Step 5 — Recover Site (Coarse-to-Fine Relocation)": "步骤 5 — 恢复站点（粗到精重定位）",
    "Step 6 — Verify Tip": "步骤 6 — 验证探针位置",
    "Recommended Workflow": "推荐工作流程",
    
    # === SECTION 8 ===
    "8.1 Artefact Detection (YOLOv8)": "8.1 伪影检测（YOLOv8）",
    "8.2 Phase-2 Machine Learning Models": "8.2 第二阶段机器学习模型",
    "Same-Site Classifier": "同站点分类器",
    "Remount Transform Predictor": "重新装载变换预测器",
    "Low-Mag Embedding Index": "低倍镜嵌入索引",
    "8.3 DeepSeek LLM Agent": "8.3 DeepSeek 大语言模型代理",
    "8.4 Training the AI Models": "8.4 训练 AI 模型",
    
    # === SECTION 9 ===
    "9.1 Automatic Data Collection": "9.1 自动数据采集",
    "9.2 Trajectory Export": "9.2 轨迹导出",
    
    # === SECTION 10 ===
    "10.1 Mouse Actions": "10.1 鼠标操作",
    "10.2 Viewport Arrow Controls": "10.2 视口方向箭头控制",
    "10.3 Button Reference": "10.3 按钮参考",
    
    # === SECTION 11 ===
    "— End of Manual —": "— 手册完 —",
}

# Full paragraph translations (long paragraphs)
long_translations = {
    "The AFM Hysteresis Simulation System is an interactive desktop application that simulates "
    "the behaviour of an Atomic Force Microscope (AFM) scanning stage. It models the Prandtl-Ishlinskii (PI) "
    "hysteresis effect inherent in piezoelectric actuators, provides AI-based inverse compensation, "
    "and implements a full sample-relocation pipeline using computer vision and machine learning.": 
    "AFM 迟滞效应模拟系统是一个交互式桌面应用程序，用于模拟原子力显微镜（AFM）扫描载物台的行为。"
    "它对压电执行器固有的 Prandtl-Ishlinskii（PI）迟滞效应进行建模，提供基于 AI 的逆向补偿，"
    "并实现了基于计算机视觉和机器学习的完整样品重定位流程。",
    
    "Real-time PI hysteresis modelling for X and Y scanner axes with configurable play operators.":
    "对 X 和 Y 扫描轴进行实时 PI 迟滞建模，支持可配置的 play 算子。",
    
    "MLP neural-network inverse model for hysteresis compensation (AI compensation mode).":
    "用于迟滞补偿的 MLP 神经网络逆向模型（AI 补偿模式）。",
    
    "Interactive optical-microscope simulation with variable zoom (0.25× to 10×), defocus blur, and camera lift.":
    "交互式光学显微镜模拟，支持可变倍率（0.25× 至 10×）、离焦模糊和相机升降。",
    
    "Mouse-driven navigation: click-to-move, right-click origin definition, and viewport panning.":
    "鼠标驱动导航：点击移动、右键定义原点以及视口平移。",
    
    "Six-step sample relocation pipeline: Save Region → Remount → Pick Origin → Find Origin → Recover Site → Verify Tip.":
    "六步样品重定位流程：保存区域 → 重新装载 → 选择原点 → 查找原点 → 恢复站点 → 验证探针。",
    
    "Phase-2 machine learning: Random Forest same-site classifier, remount-transform predictor, and low-mag embedding retrieval.":
    "第二阶段机器学习：随机森林同站点分类器、重新装载变换预测器以及低倍镜嵌入检索。",
    
    "YOLOv8 artefact detector for fiducial marks (cross, fiducial, circle, square).":
    "用于基准标记（十字、基准点、圆形、方形）的 YOLOv8 伪影检测器。",
    
    "Automatic data collection and CSV export for hysteresis-loop analysis.":
    "自动数据采集和 CSV 导出，用于迟滞回线分析。",
    
    "Dockable, resizable UI panels with persistent layout save/load.":
    "可停靠、可调整大小的 UI 面板，支持布局的持久保存和加载。",
    
    "DeepSeek LLM agent integration for conversational AFM simulation control.":
    "集成 DeepSeek 大语言模型代理，实现对话式 AFM 模拟控制。",
    
    "The project consists of 21 Python files organised into four logical layers: "
    "(1) Simulation Engine — hysteresis.py, afm_state.py, afm_optics_model.py; "
    "(2) User Interface — afm_ui.py, afm_control_panel.py, afm_callbacks.py, afm_animation.py; "
    "(3) Relocation & AI — afm_relocation.py, afm_phase2_ml.py, artefact_detector.py, image_matching.py; "
    "(4) Training Pipeline — preprocess_data.py, data_collection.py, train_*.py.":
    "该项目由 21 个 Python 文件组成，分为四个逻辑层："
    "（1）模拟引擎 — hysteresis.py、afm_state.py、afm_optics_model.py；"
    "（2）用户界面 — afm_ui.py、afm_control_panel.py、afm_callbacks.py、afm_animation.py；"
    "（3）重定位与 AI — afm_relocation.py、afm_phase2_ml.py、artefact_detector.py、image_matching.py；"
    "（4）训练管线 — preprocess_data.py、data_collection.py、train_*.py。",
    
    "Python 3.10 or later.": "Python 3.10 或更高版本。",
    "Windows, macOS, or Linux with a graphical display (matplotlib requires a GUI backend).":
    "Windows、macOS 或 Linux，需具备图形显示能力（matplotlib 需要 GUI 后端）。",
    "Approximately 2 GB free disk space for dependencies (including PyTorch for YOLOv8).":
    "约 2 GB 可用磁盘空间用于依赖项（包括 YOLOv8 所需的 PyTorch）。",
    
    "Install all dependencies with a single command:": "使用一条命令安装所有依赖：",
    
    "This installs numpy, matplotlib, opencv-python, joblib, pandas, scikit-learn, and ultralytics (which brings PyTorch). "
    "The optional python-docx package is only needed to regenerate this manual.":
    "这将安装 numpy、matplotlib、opencv-python、joblib、pandas、scikit-learn 和 ultralytics（后者会附带 PyTorch）。"
    "可选的 python-docx 包仅用于重新生成本手册。",
    
    "From the project root directory, run:": "在项目根目录下运行：",
    
    "When the application starts, it loads a default sample image (if configured) or generates a synthetic \"A\"-shaped "
    "test pattern with artefact marks. The Matplotlib window displays the main viewport and several dockable control panels.":
    "启动时，应用程序会加载默认样品图像（如果已配置），或生成带有伪影标记的合成"A"形测试图案。"
    "Matplotlib 窗口显示主视口和多个可停靠的控制面板。",
    
    "The system attempts to load the image path stored in afm_default_settings.json. "
    "If no path is configured, a procedural sample is generated. Use the \"Load Image\" button in the Utility Dock "
    "to load a custom image, or \"Load Default\" to reload the saved default.":
    "系统尝试加载 afm_default_settings.json 中存储的图像路径。"
    "如果未配置路径，则会生成一个程序化样品。使用工具面板中的"加载图像"按钮加载自定义图像，"
    "或使用"加载默认"重新加载已保存的默认图像。",
    
    "The application window is divided into a large central Viewport and seven dockable panels. "
    "Each panel can be dragged by its title bar, resized from the bottom-right corner, minimised, or restored.":
    "应用程序窗口分为一个大型中央视口和七个可停靠面板。"
    "每个面板可通过标题栏拖动、从右下角调整大小、最小化或恢复。",
}

# Process each paragraph in the document
for para in doc.paragraphs:
    text = para.text.strip()
    style_name = para.style.name
    
    # Handle empty paragraphs
    if not text:
        new_doc.add_paragraph()
        continue
    
    # Check exact translation first
    translated = translate_paragraph(None, para, None)
    if translated:
        if style_name.startswith('Heading'):
            level = int(style_name.split()[-1])
            new_doc.add_heading(translated, level=level)
        else:
            p = new_doc.add_paragraph()
            run = p.add_run(translated)
            # Copy formatting
            if para.alignment:
                p.alignment = para.alignment
            if para.runs:
                first_run = para.runs[0]
                if first_run.bold:
                    run.bold = True
                if first_run.italic:
                    run.italic = True
                if first_run.font.size:
                    run.font.size = first_run.font.size
                if first_run.font.color and first_run.font.color.rgb:
                    run.font.color.rgb = first_run.font.color.rgb
        continue
    
    # Check long translations
    if text in long_translations:
        new_doc.add_paragraph(long_translations[text])
        continue
    
    # For list items and other short paragraphs, try partial translation
    # First check if it's a known type (list bullet, list number)
    if para.style.name == 'List Bullet' or para.style.name == 'List Number':
        # Translate common patterns in list items
        new_doc.add_paragraph(text)  # Keep original for list items we haven't mapped
        continue
    
    # For unhandled paragraphs, keep the original text
    new_doc.add_paragraph(text)

# Save
output_path = 'AFM_Hysteresis_Simulation_Operation_Manual_CN.docx'
new_doc.save(output_path)
print(f'Chinese manual saved to: {output_path}')
