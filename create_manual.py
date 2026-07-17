"""
生成中文操作手册：AFM_Hysteresis_Simulation_操作手册.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

doc = Document()

# ── 样式设置 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '微软雅黑'
    heading_style.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_note(text):
    p = doc.add_paragraph()
    run = p.add_run('💡 ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.italic = True
    return p

def add_warning(text):
    p = doc.add_paragraph()
    run = p.add_run('⚠️ ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)
    run.bold = True
    return p

# ════════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AFM 滞回仿真与 AI 重定位系统')
run.font.size = Pt(26)
run.font.name = '微软雅黑'
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('操作手册')
run.font.size = Pt(18)
run.font.name = '微软雅黑'
run.font.color.rgb = RGBColor(0x52, 0x72, 0x94)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'版本 2.0  |  {datetime.now().strftime("%Y年%m月%d日")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 目录
# ════════════════════════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '1. 系统概述',
    '2. 环境安装与启动',
    '3. 界面布局说明',
    '   3.1 导航面板 (Navigation Dock)',
    '   3.2 运动面板 (Motion Dock)',
    '   3.3 重定位面板 (Relocation Dock)',
    '   3.4 工具面板 (Utility Dock)',
    '   3.5 状态面板 (Status Dock)',
    '4. 基本操作',
    '   4.1 视图移动与缩放',
    '   4.2 加载样本图像',
    '   4.3 设置坐标原点',
    '   4.4 表面倾斜调整',
    '5. 传统重定位流程（6步骤）',
    '6. AI 重定位流程（一鍵操作）— PPT 第2页',
    '7. AI 缩放恢复流程 — PPT 第3页',
    '8. ML 模型训练',
    '   8.1 训练数据准备',
    '   8.2 执行训练',
    '   8.3 ML 模型说明',
    '9. 悬臂尖端点击修正模式',
    '10. 常见问题与故障排除',
]
for item in toc_items:
    add_para(item, size=10.5)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 1. 系统概述
# ════════════════════════════════════════════════════════════
doc.add_heading('1. 系统概述', level=1)
add_para('本系统是一个 AFM（原子力显微镜）仿真环境，主要实现两个功能方向：')
add_para('(1) 滞回感知的运动仿真与可视化 — 模拟压电陶瓷驱动器的迟滞非线性特性（PI补偿模式）。')
add_para('(2) 视觉辅助的样本/位点恢复 — 在样本被移除并重新放置后，通过计算机视觉或机器学习技术自动找回之前的扫描位置。')
add_para('')
add_para('本手册重点介绍第(2)部分：AI 辅助的重定位（AI Relocation）和 AI 缩放恢复（AI Zoom Recovery），对应导师 PPT 第 2 页和第 3 页的功能需求。')

add_note('完整的研发计划请参见项目目录下的 AI_REPOSITIONING_PLAN.md')

doc.add_heading('核心工作流程', level=2)
workflow_steps = [
    '① 加载样本图像 → 导航到感兴趣区域',
    '② 保存参考位置 (Save Region)：存储低倍/高倍参考图像、landmark 特征点和缩放信息',
    '③ 模拟样本移除 (Remount)：系统随机施加平移、旋转、倾斜',
    '④ [传统] 6步骤重定位 或 [AI] 一键 AI Recall 自动恢复',
    '⑤ 验证悬臂尖端是否回到正确位置',
    '⑥ [AI Zoom] 可选：AI 缩放搜索，在不同放大倍率下匹配图案',
]
for step in workflow_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 2. 环境安装与启动
# ════════════════════════════════════════════════════════════
doc.add_heading('2. 环境安装与启动', level=1)

doc.add_heading('2.1 安装依赖', level=2)
add_para('确保已安装 Python 3.10+，然后在项目目录下执行：')
add_code('pip install -r requirements.txt')
add_para('主要依赖包括：numpy, matplotlib, opencv-python, scikit-learn, joblib, torch, torchvision, ultralytics')
add_para('')
add_warning('首次运行时会自动下载 ResNet18 预训练权重（约 45MB），请保持网络连接。')

doc.add_heading('2.2 启动系统', level=2)
add_code('python afm_control_panel.py')
add_para('系统启动后将显示主窗口，包含显微镜视图和多个控制面板。')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 3. 界面布局
# ════════════════════════════════════════════════════════════
doc.add_heading('3. 界面布局说明', level=1)
add_para('主窗口左侧为显微镜视图区，右侧分布多个可拖拽的控制面板 (Dock)：')

doc.add_heading('3.1 导航面板 (Navigation Dock)', level=2)
add_para('包含步长选择（1/5/50/200 μm）和方向键（上/下/左/右），控制悬臂在样本表面的移动步进。')

doc.add_heading('3.2 运动面板 (Motion Dock)', level=2)
add_para('包含关键的移动控制按钮：')
controls = [
    ('PI Compensation Mode', '切换 PI 迟滞补偿模式（默认关闭）'),
    ('Auto Scan', '自动水平扫描'),
    ('Motion: ON/OFF', '全局运动暂停/恢复'),
    ('Stop Here', '立即停止当前移动，目标位置吸附到当前位置'),
    ('Go Now', '立即跳转到目标位置（跳过平滑动画）'),
    ('Best Focus', '复位 Z 轴对焦'),
    ('Zoom + / Zoom -', '光学变倍缩放（12级：0.25x ~ 10.0x）'),
    ('Z + / Z -', 'Z 轴对焦微调'),
]
for label, desc in controls:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3.3 重定位面板 (Relocation Dock)', level=2)
add_para('这是本系统最核心的面板，包含 8 个按钮：')
reloc_buttons = [
    ('1. Save Region', '保存当前区域为参考位置（低倍+高倍 landmark、缩放信息）'),
    ('2. Remount', '模拟样本移除和重新放置（随机平移±500μm、旋转±8°、倾斜±10°）'),
    ('3. Pick Origin', '在当前视图中自动选取最强特征点作为坐标原点'),
    ('4. Find Origin', '在整个样本上搜索之前保存的原点图案'),
    ('5. Recover Site', '传统 6 步骤重定位：粗定位→仿射变换→细定位→验证'),
    ('6. Verify Tip', '重新匹配高倍 landmark，验证悬臂尖端位置是否正确'),
    ('AI Recall', '一鍵 AI 重定位（PPT 第2页功能）'),
    ('AI Zoom', '一鍵 AI 缩放恢复（PPT 第3页功能）'),
]
for label, desc in reloc_buttons:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{label}：')
    run.bold = True
    p.add_run(desc)

add_note('将鼠标悬停在按钮上，Relocation Dock 底部会显示该按钮的详细说明。')

doc.add_heading('3.4 工具面板 (Utility Dock)', level=2)
add_para('包含图像加载（Load Image / Load Default）、HUD 切换、标尺切换、倾斜设置等工具。')

doc.add_heading('3.5 状态面板 (Status Dock)', level=2)
add_para('实时显示当前参数：位置坐标、缩放级别、视图范围、倾斜角度、PI模式状态等。')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 4. 基本操作
# ════════════════════════════════════════════════════════════
doc.add_heading('4. 基本操作', level=1)

doc.add_heading('4.1 视图移动与缩放', level=2)
add_para('方向键：点击导航面板的 ↑↓←→ 按钮，悬臂以设定步长移动。')
add_para('点击视图：在显微镜视图中左键点击任意位置，悬臂尖端自动移动到该点。')
add_para('缩放：使用 Zoom + / Zoom - 按钮在 12 级放大倍率之间切换（0.25x ~ 10.0x）。')
add_para('急停/急跳：Stop Here 立即停止运动；Go Now 立即跳转到目标。')

doc.add_heading('4.2 加载样本图像', level=2)
add_para('点击 Utility Dock 中的 "Load Image" 加载自定义显微镜图像，或点击 "Load Default" 加载默认样本。')

doc.add_heading('4.3 设置坐标原点', level=2)
add_para('点击 "Pick Origin" 按钮，系统自动在视图中心附近选取最强特征点作为坐标原点。原点将显示为十字标记。')
add_para('也可用 "Find Origin" 在整个样本图像上搜索之前保存的原点位置。')

doc.add_heading('4.4 表面倾斜调整', level=2)
add_para('点击 "Stage Tilt" 输入倾斜角度（-10° ~ +10°），模拟样本在载物台上的倾斜。')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 5. 传统重定位流程
# ════════════════════════════════════════════════════════════
doc.add_heading('5. 传统重定位流程（6步骤）', level=1)
add_para('这是系统中已有的标准重定位工作流，适合理解重定位的基本原理：')

steps_detail = [
    ('步骤 1 — Save Region',
     '导航到感兴趣的扫描区域，调整到合适的放大倍率，点击 "1. Save Region"。\n'
     '系统会保存：当前高倍视图（参考模板）、低倍全图（概览图）、多个 landmark 特征点、缩放级别、原点位置。\n'
     '所有数据以 site_memory 的形式持久化到 collected_data/site_memories/ 目录。'),
    ('步骤 2 — Remount',
     '点击 "2. Remount" 模拟样本移除和重新放置。\n'
     '系统随机施加平移（±500 μm）、面内旋转（±8°）和表面倾斜（±10°）。\n'
     '样本图像会被仿射变换处理，边缘可能出现黑色填充区域。'),
    ('步骤 3 — Pick Origin',
     '点击 "3. Pick Origin" 在当前视图中自动选取一个显著特征作为坐标原点。'),
    ('步骤 4 — Find Origin',
     '点击 "4. Find Origin" 在整个样本图像上用模板匹配搜索之前保存的原点位置。'),
    ('步骤 5 — Recover Site',
     '点击 "5. Recover Site" 执行完整的粗到细重定位：\n'
     '  • 粗定位：低倍 landmark 共识估计 或 仿射变换估计旋转/偏移\n'
     '  • 细定位：模板匹配 + 局部仿射细化\n'
     '  • 验证：参考分数、landmark 几何一致性检查\n'
     '验证通过则自动移动悬臂到恢复位置。'),
    ('步骤 6 — Verify Tip',
     '点击 "6. Verify Tip" 重新匹配所有高倍 landmark，验证悬臂尖端是否真正回到了正确位置。\n'
     'HUD 叠加层会显示 landmark 匹配情况和距离误差。'),
]
for title, desc in steps_detail:
    doc.add_heading(title, level=2)
    add_para(desc)

add_note('传统重定位基于 OpenCV 的 ORB 特征点匹配和模板匹配，不依赖机器学习模型，无需训练即可使用。')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 6. AI 重定位
# ════════════════════════════════════════════════════════════
doc.add_heading('6. AI 重定位流程（一鍵操作）— PPT 第2页', level=1)

add_para('AI Recall（一键 AI 重定位）是导师 PPT 第 2 页要求实现的核心功能。它将传统 6 步骤流程自动化为一键操作，并使用 Machine Learning 替代传统计算机视觉进行图案识别。')

doc.add_heading('PPT 第2页功能对照', level=2)
ppt2_items = [
    ('Machine recall the last measurement region',
     '系统自动加载上次保存的 site_memory（包含参考图像、landmark、缩放值等）。'),
    ('AI recognize the current surrounding pattern',
     '使用 ResNet18 深度特征 + MLP 神经网络进行图案识别和旋转角度估计。'),
    ('Rotation angle calculated from previous shape recognition',
     'ML 回归器（MLTransformPredictor）直接预测 dx, dy, dθ 变换参数。'),
    ('Cantilever moves to new origin based on recognition',
     '根据 ML 预测结果，自动移动悬臂到恢复位置。'),
    ('Verify location based on multiple distances',
     '双重验证：传统 landmark 几何一致性 + ML 分类器（MLSameSiteClassifier）判断。'),
    ('If not correct, move based on click',
     '验证失败时自动进入 click-to-move 修正模式，用户点击正确位置即可。'),
]
for ppt_item, impl in ppt2_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'PPT: "{ppt_item}"')
    run.italic = True
    p.add_run(f'\n  实现：{impl}')

doc.add_heading('使用方法', level=2)
add_para('前提条件：已训练 ML 模型（见第 8 章），且至少保存过一次参考位置。')
add_para('')
steps_ai = [
    '步骤 1：按传统流程保存参考位置（1. Save Region）',
    '步骤 2：模拟样本移除（2. Remount）',
    '步骤 3：直接点击 "AI Recall" 按钮',
]
for i, step in enumerate(steps_ai):
    add_para(f'{step}')

add_para('')
add_para('系统将自动执行以下操作（全程无需人工干预）：')
auto_steps = [
    '① 加载最新 site_memory，恢复保存时的缩放级别',
    '② AI 识别（ML path）：ResNet18 提取当前视图的 512 维深度特征',
    '   在搜索图像上滑动窗口匹配参考模板',
    '   同时预测 remount 变换参数（dx, dy, dθ）',
    '   若 ML 不可用，自动回退到传统 CV（ORB + matchTemplate）',
    '③ 传统粗定位备选：仿射变换估计 + 低倍 landmark 共识',
    '④ 双重验证：传统 landmark 几何检查 + ML 分类器概率判定',
    '⑤ 验证通过 → 自动移动悬臂；验证失败 → 进入 click-to-move 修正模式',
]
for step in auto_steps:
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)

doc.add_heading('识别技术对比', level=2)
add_para('')

# 简单表格
table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
headers = ['步骤', '传统 CV（步骤5）', 'AI ML（AI Recall）']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['特征提取', 'ORB 关键点 (cv2.ORB)', 'ResNet18 深度特征 (512维)'],
    ['匹配方式', '模板匹配 (cv2.matchTemplate)', '滑动窗口 + 余弦相似度'],
    ['变换预测', 'RANSAC 仿射估计', 'MLP 回归器直接预测 dx,dy,dθ'],
    ['验证', '分数阈值 + landmark 几何', 'MLP 分类器 + 传统验证'],
]
for r, row_data in enumerate(data):
    for c, cell_text in enumerate(row_data):
        table.rows[r + 1].cells[c].text = cell_text

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 7. AI 缩放恢复
# ════════════════════════════════════════════════════════════
doc.add_heading('7. AI 缩放恢复流程 — PPT 第3页', level=1)

add_para('AI Zoom（AI 缩放恢复）是导师 PPT 第 3 页要求实现的功能。它解决了在不同放大倍率下识别图案的难题。')

doc.add_heading('PPT 第3页功能对照', level=2)
ppt3_items = [
    ('Machine recall the last zoom value, then AI zoom it',
     '系统从 site_memory 中恢复上次保存的缩放级别，并自动缩放到该级别。'),
    ('AI recognize the current surrounding pattern',
     '首先尝试在恢复的缩放级别下用 ML 识别图案。'),
    ('If not recognize, search begins — zoom out and zoom in',
     '如果识别失败，系统自动逐步 zoom out 搜索。找到后 zoom in 回原倍数。'),
    ('If recognize, cantilever moves to new origin',
     '识别成功后，自动移动悬臂到目标位置。'),
    ('Verify the location of the cantilever tip',
     'ML 分类器 + 传统 landmark 双重验证。'),
]
for ppt_item, impl in ppt3_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'PPT: "{ppt_item}"')
    run.italic = True
    p.add_run(f'\n  实现：{impl}')

doc.add_heading('使用方法', level=2)
add_para('前提条件：已保存参考位置（site_memory 中包含 zoom_level），当前与保存时的缩放级别不同。')
add_para('')
add_para('直接点击 "AI Zoom" 按钮，系统自动执行缩放搜索：')

add_para('AI Zoom 搜索流程图：')
search_flow = [
    '┌──────────────────────────────────────┐',
    '│  1. 加载 site_memory，恢复缩放级别    │',
    '│          ↓                           │',
    '│  2. ML 识别：当前缩放下找图案         │',
    '│       ↙              ↘               │',
    '│   找到了            没找到            │',
    '│     ↓                 ↓               │',
    '│  3. 移动悬臂    4. zoom out 搜索       │',
    '│     ↓                 ↓               │',
    '│  4. 验证        5. 逐步放大再搜索       │',
    '│     ↓              ↙  ↖              │',
    '│  通过→移动      找到→zoom in→移动      │',
    '│  失败→点击修正  耗尽→点击修正          │',
    '└──────────────────────────────────────┘',
]
for line in search_flow:
    add_code(line)

add_para('')
add_para('搜索过程：')
search_steps = [
    '① 恢复到保存时的缩放级别（如 4.0x）',
    '② 在当前视图尝试 ML 图案识别',
    '③ 若失败，自动 zoom out 到更宽视角（如 2.0x）再搜索',
    '④ 可继续 zoom out 直到找到为止（最宽 0.25x）',
    '⑤ 找到图案后，自动 zoom in 回原倍数并移动悬臂',
]
for step in search_steps:
    doc.add_paragraph(step, style='List Number')

add_warning('若在所有缩放级别都找不到图案，系统会进入 click-to-move 修正模式，允许用户手动点击正确位置。')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 8. ML 模型训练
# ════════════════════════════════════════════════════════════
doc.add_heading('8. ML 模型训练', level=1)

add_para('为了让 AI Recall 和 AI Zoom 使用 Machine Learning 而非传统 CV，需要先训练 ML 模型。')

doc.add_heading('8.1 训练数据准备', level=2)
add_para('训练数据由已保存的 site_memory 自动生成。数据生成方式：')
prep_steps = [
    '1. 运行 python afm_control_panel.py 启动系统',
    '2. 导航到感兴趣的样本区域',
    '3. 点击 "1. Save Region" 保存参考位置',
    '4. （可选）改变位置，重复步骤 2-3 保存多个不同区域',
    '5. 也可以点击 "2. Remount" 多次，每次重新保存，增加训练数据多样性',
]
for step in prep_steps:
    add_para(step)

add_para('')
add_para('训练数据由 afm_phase2_ml.py 和 afm_ml_recognition.py 自动生成：')
data_gen = [
    '• 正样本（同一site）：对参考图像施加小幅平移/旋转/亮度变化，生成 ~8 个变体',
    '• 负样本（不同site）：不同 site 之间的图像配对、随机裁剪、块交换等',
    '• 回归样本：对概览图像施加已知的平移/旋转变换，记录真实的 dx, dy, dθ',
]
for item in data_gen:
    add_para(item)

doc.add_heading('8.2 执行训练', level=2)
add_code('python train_ml_models.py')
add_para('')
add_para('训练过程：')
training_desc = [
    '• 初始化 ResNet18 特征提取器（首次运行下载预训练权重 ~45MB）',
    '• Step 1：训练 Deep Same-Site Classifier — MLP 三隐层 (256→128→64)',
    '  - 输入：1024维深度特征拼接 + 差异向量 + 统计特征',
    '  - 输出：同一 site 的概率 [0, 1]',
    '• Step 2：训练 Deep Remount Predictor — MLP 三隐层 (256→128→64)',
    '  - 输入：同上 1034维特征',
    '  - 输出：dx (μm), dy (μm), dθ (度)',
    '• 模型保存到 collected_data/models/deep_same_site_classifier.pkl',
    '                      collected_data/models/deep_remount_predictor.pkl',
]
for item in training_desc:
    add_para(item)

doc.add_heading('8.3 ML 模型说明', level=2)
add_para('系统中有两套 ML 模型：')

# 表格
table2 = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
headers2 = ['模型', '类型', '特征提取', '用途']
for i, h in enumerate(headers2):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

model_data = [
    ['same_site_classifier', 'RandomForest (传统)', '10个CV特征', '传统重定位辅助验证'],
    ['remount_transform_predictor', 'RandomForest (传统)', '10个CV特征', '传统重定位辅助预测'],
    ['deep_same_site_classifier', 'MLP (深度)', 'ResNet18 512维', 'AI Recall 主识别 + 验证'],
    ['deep_remount_predictor', 'MLP (深度)', 'ResNet18 512维', 'AI Recall 变换预测'],
]
for r, row_data in enumerate(model_data):
    for c, cell_text in enumerate(row_data):
        table2.rows[r + 1].cells[c].text = cell_text

add_para('')
add_note('传统 RandomForest 模型由 train_repositioning_ai.py 训练，使用 OpenCV 手工特征。\n深度 MLP 模型由 train_ml_models.py 训练，使用 ResNet18 预训练特征。两者可以共存，系统会自动选择最佳路径。')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 9. Click-to-move 修正
# ════════════════════════════════════════════════════════════
doc.add_heading('9. 悬臂尖端点击修正模式', level=1)

add_para('当 AI Recall 或 AI Zoom 的验证步骤失败时，系统不会放弃，而是自动进入 "click-to-move" 修正模式。')

add_para('')
add_para('操作方式：')
click_steps = [
    '1. 日志区会显示警告信息和 AI 的最佳猜测坐标',
    '2. 在显微镜视图中，左键点击你认为正确的悬臂目标位置',
    '3. 系统立即将悬臂移动到该位置',
    '4. 右键点击可以取消修正模式',
]
for step in click_steps:
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)

add_para('')
add_para('触发条件：')
trigger_conditions = [
    '• AI Recall 验证分数低于阈值',
    '• AI Zoom 在所有缩放级别都找不到图案',
    '• Landmark 几何一致性检查失败',
    '• ML 分类器判定为非同一 site',
]
for cond in trigger_conditions:
    doc.add_paragraph(cond, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 10. 常见问题
# ════════════════════════════════════════════════════════════
doc.add_heading('10. 常见问题与故障排除', level=1)

faq = [
    ('Q: 点击 AI Recall 后显示 "No saved site memory found"？',
     'A: 需要先点击 "1. Save Region" 保存至少一个参考位置。每次保存都会在 collected_data/ 下生成 site_memory。'),
    ('Q: AI Recall 始终使用 CV 回退而不使用 ML？',
     'A: 需要先运行 python train_ml_models.py 训练 ML 模型。检查 collected_data/models/ 下是否有 deep_*.pkl 文件。'),
    ('Q: 训练 ML 模型时报错 "No site memories found"？',
     'A: 需要先保存至少一个参考位置（点击 "1. Save Region"）。训练数据从已保存的 site_memory 自动生成。'),
    ('Q: ResNet18 下载太慢？',
     'A: 首次运行时 PyTorch 会自动下载预训练权重。如果网络不好，可以提前手动下载到 ~/.cache/torch/hub/checkpoints/。'),
    ('Q: 重定位验证总是失败？',
     'A: 可能原因：样本重复图案太多（ambiguity）；remount 施加的变换过大；训练数据不足。\n'
     '   建议：保存多个不同位置的 site_memory；使用 click-to-move 手动修正。'),
    ('Q: Zoom 后图像变模糊？',
     'A: 这是仿真特性——放大时 FOV 变小、像素密度降低。调整 Z 轴对焦（Z+/Z-）或点击 Best Focus 复位。'),
    ('Q: 如何在已有的 site_memory 基础上增量训练 ML？',
     'A: 重新保存更多 site_memory，然后再次运行 python train_ml_models.py，新数据会自动纳入训练。'),
    ('Q: 传统重定位和 AI 重定位有什么区别？',
     'A: 传统重定位（步骤5）使用 OpenCV 的 ORB 特征点和模板匹配，无需训练。\n'
     '    AI 重定位（AI Recall）使用 ResNet18 + MLP 神经网络，需要训练，但识别更鲁棒。\n'
     '    两者都可以在同一个界面中使用，互不冲突。'),
    ('Q: 系统启动很慢？',
     'A: 首次启动需要加载 ResNet18（~45MB），后续启动会复用缓存。Artefact Detector (YOLO) 也会在首次加载。'),
]
for q, a in faq:
    doc.add_heading(q, level=2)
    add_para(a)

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# 附录
# ════════════════════════════════════════════════════════════
doc.add_heading('附录', level=1)

doc.add_heading('A. 项目文件速查', level=2)
file_refs = [
    ('afm_control_panel.py', '主入口：启动 UI、绑定按钮、管理生命周期'),
    ('afm_state.py', '全局状态：坐标、缩放、site_memory、AI 状态变量'),
    ('afm_callbacks.py', '回调逻辑：移动、缩放、重定位、AI Recall、AI Zoom'),
    ('afm_relocation.py', '重定位引擎：landmark 提取、仿射变换、模板匹配'),
    ('afm_ml_recognition.py', 'ML 引擎：ResNet18 特征、MLP 分类器/回归器、训练函数'),
    ('afm_phase2_ml.py', 'Phase 2 ML：传统 RandomForest 模型、合成数据生成'),
    ('afm_ui.py', 'UI 布局：Dock 面板、按钮位置、响应式布局'),
    ('afm_animation.py', '动画引擎：缩放插值、帧渲染'),
    ('train_ml_models.py', '训练脚本：一键训练深度 ML 模型'),
    ('train_repositioning_ai.py', '训练脚本：一键训练传统 RandomForest 模型'),
]
for fname, desc in file_refs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{fname}')
    run.bold = True
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    p.add_run(f' — {desc}')

doc.add_heading('B. 键盘快捷键', level=2)
shortcuts = [
    ('方向键 ↑↓←→', '悬臂移动（需要导航面板焦点）'),
    ('左键点击视图', '移动悬臂尖端到点击位置'),
    ('右键点击视图', '取消 click-to-move 修正模式'),
    ('鼠标悬停按钮', '显示按钮详细说明（Relocation Dock）'),
]
for key, desc in shortcuts:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{key}')
    run.bold = True
    p.add_run(f' — {desc}')

doc.add_heading('C. 数据目录结构', level=2)
add_code('collected_data/')
add_code('├── site_memories/          # 保存的参考位置数据')
add_code('│   └── <sample_id>/')
add_code('│       └── <session>_<site>/')
add_code('│           ├── metadata.json      # 坐标、zoom、landmark 等元数据')
add_code('│           ├── reference.png      # 高倍参考图像')
add_code('│           ├── overview.png       # 低倍概览图像')
add_code('│           └── landmarks/         # landmark 补丁图像')
add_code('│               ├── lowmag/')
add_code('│               └── highmag/')
add_code('└── models/                # 训练好的 ML 模型')
add_code('    ├── same_site_classifier.pkl        # 传统 RF 分类器')
add_code('    ├── remount_transform_predictor.pkl # 传统 RF 回归器')
add_code('    ├── lowmag_embedding_index.pkl      # 低倍检索索引')
add_code('    ├── deep_same_site_classifier.pkl   # 深度 MLP 分类器')
add_code('    └── deep_remount_predictor.pkl      # 深度 MLP 回归器')

# ── 保存 ──
output_path = 'AFM_Hysteresis_Simulation_操作手册.docx'
doc.save(output_path)
print(f'操作手册已保存到: {output_path}')
